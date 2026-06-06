# -*- coding: utf-8 -*-
"""按参考模板格式生成作业三 Word 报告。"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
REPORT_JSON_ZH = os.path.join(OUTPUT_DIR, "chinese_ai_report.json")
REPORT_JSON_EN = os.path.join(OUTPUT_DIR, "newsgroups_report.json")
SUMMARY_JSON = os.path.join(OUTPUT_DIR, "summary.json")

OUT_DESKTOP = r"d:\HuaweiMoveData\Users\xiaos\Desktop\作业3-陈喜盈(3223004641).docx"
OUT_LOCAL = os.path.join(BASE_DIR, "作业3-陈喜盈(3223004641).docx")


def set_run_font(run, name="宋体", size_pt=12, bold=None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name if name != "黑体" else "Times New Roman")
    rFonts.set(qn("w:hAnsi"), name if name != "黑体" else "Times New Roman")
    rFonts.set(qn("w:eastAsia"), name)


def add_empty_lines(doc, n=1):
    for _ in range(n):
        doc.add_paragraph("")


def add_para(
    doc,
    text,
    *,
    bold=False,
    size=12,
    font="宋体",
    align=None,
    space_before=None,
    space_after=6,
    first_indent=None,
    line_spacing=1.15,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    run = p.add_run(text)
    set_run_font(run, font, size, bold=bold)
    return p


def add_body(doc, text, space_after=6):
    return add_para(doc, text, size=12, font="宋体", space_after=space_after)


def add_h1(doc, text):
    return add_para(doc, text, bold=True, size=22, font="宋体", space_before=24, space_after=0)


def add_h2(doc, text):
    return add_para(doc, text, bold=True, size=16, font="黑体", space_before=10, space_after=0)


def add_code(doc, text):
    return add_para(doc, text, size=9, font="Times New Roman", space_after=6)


def add_figure(doc, caption, image_path, width_cm=14):
    add_body(doc, caption)
    if os.path.isfile(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        add_body(doc, f"[图片缺失: {image_path}]")
    add_empty_lines(doc, 1)


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def topic_lines(topics, n=6):
    lines = []
    for t in topics[:5]:
        words = "、".join(w for w, _ in t["terms"][:n])
        lines.append(f"  Topic {t['topic_id']}：{words}")
    return "\n".join(lines)


def build_document():
    zh = load_json(REPORT_JSON_ZH)
    en = load_json(REPORT_JSON_EN)
    summary = load_json(SUMMARY_JSON)

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    # 封面
    add_empty_lines(doc, 5)
    add_para(
        doc,
        "文本信息处理 作业 ",
        bold=True,
        size=26,
        font="宋体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )
    add_empty_lines(doc, 2)
    add_para(
        doc,
        "学生学院             计算机学院             ",
        size=16,
        font="黑体",
        first_indent=Cm(2),
        space_after=0,
    )
    add_para(
        doc,
        "专业班级            软件工程23(4)          ",
        size=16,
        font="黑体",
        first_indent=Cm(2),
        space_after=0,
    )
    add_para(
        doc,
        "学生姓名、学号     陈喜盈(3223004641)               ",
        size=16,
        font="黑体",
        first_indent=Cm(2),
        space_after=0,
    )
    add_empty_lines(doc, 3)
    add_para(doc, "2026  年 6  月  ", size=16, font="宋体", space_after=0)
    add_empty_lines(doc, 5)

    add_para(
        doc,
        "作业三　文档的隐含义/含义/主题挖掘",
        bold=True,
        size=18,
        font="宋体",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_empty_lines(doc, 1)

    # 一、作业要求
    add_h1(doc, "一、作业要求")
    reqs = [
        "准备一份作业3，收集、整理、产生一份个人（文本）推荐、可视化等算法。",
        "内容、方向、体裁、格式、创意不限；推荐使用作业1、2中的数据内容，并进行展示。",
        "核心流程：数据集处理 → 数据降维（LSA/LSI/LDA）→ 降维后表示 → KNN 推荐 + 可视化（PCA/t-SNE/UMAP）。",
        "加分/得分/扣分项：",
        "有趣、好的创意、数据；",
        "有新意、有效果、有脑洞的新词评估机制；",
        "生动的表达方式，图片、图表；",
        "Matplotlib、Jupyter notebook；",
        "正式的书写格式，参考毕业设计格式；",
        "个人作业；",
        "不允许抄袭、参考要列出出处，并作为文献参考。",
    ]
    for r in reqs:
        add_para(doc, r, size=10.5, font="宋体", space_after=3)

    # 二、作业内容介绍
    add_h1(doc, "二、作业内容介绍、实现思路等")
    add_h2(doc, "1. 选题背景")
    add_body(
        doc,
        "文本挖掘中，降维（Dimension Reduction）是提取文档隐含义与主题结构的常用手段。"
        "LSA/LSI 基于 SVD 将高维词项-文档矩阵映射到低维语义空间；LDA 假设文档由若干隐主题混合生成；"
        "Word2vec 等方法亦可得到分布式表示。本作业以 LSA/LSI/LDA 为出发点，"
        "设计一套“降维—嵌入—推荐—可视化”框架，并复用作业一、作业二的人工智能语料进行实验展示。",
    )

    add_h2(doc, "2. 降维与推荐方案")
    add_body(
        doc,
        "中文语料来自 homework_1/ai_corpus.txt 与 homework_2 混合语料（共 75 篇文档）；"
        "英文对照实验使用 sklearn 20 Newsgroups 四个新闻组子集（120 篇）。"
        "经结巴分词与 TF-IDF 向量化后，分别训练 LSI（TruncatedSVD）与 LDA 主题模型；"
        "以 LSI 文档嵌入为特征，建立余弦距离 KNN 索引实现相似文档推荐，并辅以关键词检索。",
    )

    add_h2(doc, "3. 技术栈")
    stack = [
        "Python 3 — 编程语言；",
        "jieba — 中文分词，加载作业二 AI 用户词典；",
        "NumPy / SciPy — SVD 理论演示与矩阵运算；",
        "scikit-learn — TF-IDF、TruncatedSVD(LSI)、LDA、PCA、t-SNE、KNN；",
        "umap-learn — UMAP 非线性降维可视化；",
        "Matplotlib — 奇异值谱、主题词、嵌入散点图；",
        "Jupyter Notebook — 交互式实验与结果展示。",
    ]
    for s in stack:
        add_body(doc, s, space_after=3)

    add_h2(doc, "4. 实现流程")
    add_body(
        doc,
        "整体流程：加载语料 → 分词与 TF-IDF/BM25 → numpy SVD 演示 → LSI/LDA 训练 → "
        "文档嵌入 PCA/t-SNE/UMAP 可视化 → KNN 相似文档推荐 → 输出图表与 JSON 报告。",
    )
    add_body(
        doc,
        "说明：作业文档推荐 Gensim 框架；当前 Python 3.14 环境暂无 Gensim 预编译包，"
        "本实现使用 sklearn + numpy 完成等价流程，算法思想与作业要求一致。",
    )

    # 三、运行结果
    add_h1(doc, "三、运行结果与测试")
    add_h2(doc, "1、代码实现")
    add_body(doc, "分词与向量化：对中文文档结巴分词后构建 TF-IDF 稀疏矩阵与词频矩阵。")
    add_code(
        doc,
        "tokenized = tokenize_docs(docs, lang='zh')\n"
        "tfidf, tfidf_vec = build_tfidf(tokenized)\n"
        "count, count_vec = build_count(tokenized)",
    )
    add_body(doc, "主题模型与推荐：LSI 得到文档嵌入，KNN 检索语义最近邻文档。")
    add_code(
        doc,
        "lsi_model, lsi_topics = train_lsi(tfidf, n_components=5)\n"
        "nn = build_knn_index(lsi_topics, metric='cosine', n_neighbors=6)\n"
        "recs = recommend(nn, lsi_topics, query_idx=0, k=5)",
    )

    add_h2(doc, "2、运行结果")
    add_body(doc, "（1）数据集规模与 SVD 能量占比：")
    add_body(
        doc,
        f"中文 AI 语料文档数：{summary['chinese_ai']['n_docs']}，前 5 维 SVD 能量占比：{summary['chinese_ai']['svd_energy']:.2%}\n"
        f"20 Newsgroups 文档数：{summary['newsgroups']['n_docs']}，前 5 维 SVD 能量占比：{summary['newsgroups']['svd_energy']:.2%}",
    )

    add_body(doc, "（2）LSI 主题词（中文语料）：")
    add_body(doc, topic_lines(zh["lsi_topics"]))

    add_body(doc, "（3）KNN 推荐示例（查询文档 #0，人工智能定义段落）：")
    rec_text = []
    for r in zh["knn_recommendations"]:
        rec_text.append(
            f"  推荐 #{r['doc_id']}（距离={r['distance']:.4f}）\n    {r['preview']}"
        )
    add_body(doc, "\n".join(rec_text))

    add_body(doc, "（4）关键词检索示例（人工智能、模型）：")
    kw_text = []
    for h in zh["keyword_hits"][:3]:
        kw_text.append(f"  #{h['doc_id']} score={h['score']:.4f} | {h['preview']}")
    add_body(doc, "\n".join(kw_text))

    add_h2(doc, "3、可视化结果")
    figures = [
        ("图 1：SVD 奇异值谱（LSA/LSI 理论基础）", "chinese_ai_svd_spectrum.png"),
        ("图 2：LSI 主题高频词分布", "chinese_ai_lsi_topics.png"),
        ("图 3：LDA 主题高频词分布", "chinese_ai_lda_topics.png"),
        ("图 4：中文语料 LSI 嵌入 — PCA 可视化", "chinese_ai_embed_pca.png"),
        ("图 5：中文语料 LSI 嵌入 — t-SNE 可视化", "chinese_ai_embed_tsne.png"),
        ("图 6：中文语料 LSI 嵌入 — UMAP 可视化", "chinese_ai_embed_umap.png"),
        ("图 7：20 Newsgroups 语料 LSI 嵌入 — UMAP 对照", "newsgroups_embed_umap.png"),
    ]
    for caption, fname in figures:
        add_figure(doc, caption, os.path.join(OUTPUT_DIR, fname))

    # 四、总结
    add_h1(doc, "四、总结")
    add_body(
        doc,
        "本作业基于作业一、作业二的 AI 语料，完成了从文本预处理、TF-IDF 向量化、"
        "LSI/LDA 主题建模到文档嵌入可视化与 KNN 推荐的完整流水线。"
        "实验表明：SVD/LSI 能有效压缩文档表示并保留主要语义结构；"
        "对“人工智能定义”类查询文档，KNN 可召回语义相近的定义、技术综述与应用段落；"
        "PCA、t-SNE、UMAP 从不同角度展示文档在隐空间中的聚类关系。"
        "后续可引入 Gensim 与中文维基百科大规模语料，进一步对比 LDA 与 LSI 在推荐任务上的效果。",
    )

    # 五、参考文献
    add_h1(doc, "五、参考文献")
    refs = [
        "[1] Gensim, Topic Modelling for humans, https://radimrehurek.com/gensim/, Accessed 2021-03-30.",
        "[2] Annoy (Approximate Nearest Neighbors Oh Yeah), https://github.com/spotify/annoy, Accessed 2021-03-30.",
        "[3] UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction, https://umap-learn.readthedocs.io/en/latest/.",
        "[4] scikit-learn User Guide — Decomposition and Manifold learning, https://scikit-learn.org/.",
        "[5] 结巴中文分词, https://github.com/fxsjy/jieba.",
        "[6] Matplotlib 官方文档, https://matplotlib.org/stable/.",
    ]
    for ref in refs:
        add_body(doc, ref, space_after=3)

    return doc


def main():
    if not os.path.isfile(REPORT_JSON_ZH):
        raise FileNotFoundError("请先运行 python main.py 生成 output 结果")

    doc = build_document()
    doc.save(OUT_LOCAL)
    print("已保存:", OUT_LOCAL)

    try:
        doc.save(OUT_DESKTOP)
        print("已保存:", OUT_DESKTOP)
    except OSError as exc:
        print("桌面保存跳过:", exc)


if __name__ == "__main__":
    main()
