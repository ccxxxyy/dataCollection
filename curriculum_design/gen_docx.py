# -*- coding: utf-8 -*-
"""按作业三模板格式生成课程设计 Word 报告。"""
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
REPORT_JSON = os.path.join(OUTPUT_DIR, "curriculum_report.json")
SUMMARY_JSON = os.path.join(OUTPUT_DIR, "summary.json")

OUT_LOCAL = os.path.join(BASE_DIR, "课程设计-陈喜盈(3223004641).docx")


def set_run_font(run, name="宋体", size_pt=12, bold=None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name if name != "黑体" else "Times New Roman")
    rFonts.set(qn("w:hAnsi"), name if name != "黑体" else "Times New Roman")
    rFonts.set(qn("w:eastAsia"), name)


def add_empty_lines(doc, n=1):
    for _ in range(n):
        doc.add_paragraph("")


def add_para(
    doc, text, *, bold=False, size=12, font="宋体", align=None,
    space_before=None, space_after=6, first_indent=None, line_spacing=1.15,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    run = p.add_run(text)
    set_run_font(run, font, size, bold=bold)
    return p


def add_body(doc, text, space_after=6):
    return add_para(doc, text, size=12, font="宋体", space_after=space_after)


def add_h1(doc, text):
    return add_para(doc, text, bold=True, size=22, font="宋体", space_before=24, space_after=0)


def add_h2(doc, text):
    return add_para(doc, text, bold=True, size=16, font="黑体", space_before=10, space_after=0)


def add_code(doc, text):
    return add_para(doc, text, size=9, font="Times New Roman", space_after=6)


def add_figure(doc, caption, image_path, width_cm=14):
    add_body(doc, caption)
    if os.path.isfile(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
    else:
        add_body(doc, f"[图片缺失: {image_path}]")
    add_empty_lines(doc, 1)


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def topic_lines(topics, n=6):
    lines = []
    for t in topics[:5]:
        words = "、".join(w for w, _ in t["terms"][:n])
        lines.append(f"  Topic {t['topic_id']}：{words}")
    return "\n".join(lines)


def newword_lines(scores, n=10):
    lines = []
    for s in scores[:n]:
        lines.append(
            f"  {s['word']}：爆发力={s['burst_index']:.4f} "
            f"(TF={s['tf']}, DF={s['df']}, 新颖度={s['novelty']})"
        )
    return "\n".join(lines)


def build_document():
    report = load_json(REPORT_JSON)
    summary = load_json(SUMMARY_JSON)
    zh = report["chinese_ai"]
    integrated = report["integrated"]
    stats = integrated["hw2_dict_stats"]

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    # 封面
    add_empty_lines(doc, 5)
    add_para(
        doc, "文本信息挖掘概论 课程设计 ",
        bold=True, size=26, font="宋体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0,
    )
    add_empty_lines(doc, 2)
    add_para(doc, "学生学院             计算机学院             ", size=16, font="黑体", first_indent=Cm(2), space_after=0)
    add_para(doc, "专业班级            软件工程23(4)          ", size=16, font="黑体", first_indent=Cm(2), space_after=0)
    add_para(doc, "学生姓名、学号     陈喜盈(3223004641)               ", size=16, font="黑体", first_indent=Cm(2), space_after=0)
    add_empty_lines(doc, 3)
    add_para(doc, "2026  年 6  月  ", size=16, font="宋体", space_after=0)
    add_empty_lines(doc, 5)

    add_para(
        doc, "课程设计　基于作业一至三的人工智能文本挖掘综合实践",
        bold=True, size=18, font="宋体", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
    )
    add_empty_lines(doc, 1)

    # 一、课程设计要求
    add_h1(doc, "一、课程设计要求")
    reqs = [
        "课程设计完全放飞自我：课题自拟、数据自拟，使用论文中的方法、模型处理分析自拟数据。",
        "推荐使用作业1、2、3中的数据内容，并进行高亮展示。",
        "核心能力展示：学习能力、包装能力、搜索定位能力、美化能力与创新脑回路。",
        "加分项：有趣、好的创意、数据；有新意、有效果、有脑洞的新词评估机制；",
        "生动的表达方式，图片、图表；Matplotlib、Jupyter notebook；",
        "正式的书写格式，参考毕业设计格式；个人作业；不允许抄袭，参考要列出出处。",
    ]
    for r in reqs:
        add_para(doc, r, size=10.5, font="宋体", space_after=3)

    # 二、课程设计内容介绍
    add_h1(doc, "二、课程设计内容介绍、实现思路等")
    add_h2(doc, "1. 选题背景")
    add_body(
        doc,
        "本课程设计以「人工智能」为主题，将作业一至三的实验成果串联为一条完整的文本挖掘流水线："
        "作业一完成语料采集与词频/TF-IDF 统计分析；作业二基于 Word2Vec+Annoy 迭代扩展定制词典并过滤新词；"
        "作业三基于 LSA/LSI/LDA 降维实现文档主题挖掘、KNN 语义推荐与 PCA/t-SNE/UMAP 可视化。"
        "课程设计在此基础上提出「新词爆发力指数」评估机制，并对三份作业数据进行综合展示与分析。",
    )

    add_h2(doc, "2. 数据与方法")
    add_body(
        doc,
        f"中文语料：作业一 ai_corpus.txt（{integrated['hw1_corpus_lines']} 条）"
        f"+ 作业二混合语料，共 {integrated['n_chinese_docs']} 篇文档；"
        f"英文对照：20 Newsgroups 四组子集 {integrated['n_english_docs']} 篇。"
        f"作业二 AI 词典从 {stats.get('AI 种子词数', '31')} 个种子词扩展至 "
        f"{stats.get('AI 最终词典词数', '223')} 个词条。",
    )
    add_body(
        doc,
        "新词爆发力指数 = 0.4×TF归一化 + 0.3×DF归一化 + 0.3×语义新颖度（相对种子词的编辑距离），"
        "用于量化作业二扩展词典中新词在语料中的传播强度与创新程度。",
    )

    add_h2(doc, "3. 技术栈")
    stack = [
        "Python 3 — 编程语言；",
        "jieba — 中文分词，加载作业二 AI 用户词典；",
        "NumPy / SciPy — SVD 理论演示与矩阵运算；",
        "scikit-learn — TF-IDF、TruncatedSVD(LSI)、LDA、PCA、t-SNE、KNN；",
        "umap-learn — UMAP 非线性降维可视化；",
        "Matplotlib — 词频、新词评估、主题词、嵌入散点图；",
        "Jupyter Notebook — 交互式实验与结果展示。",
    ]
    for s in stack:
        add_body(doc, s, space_after=3)

    add_h2(doc, "4. 实现流程")
    add_body(
        doc,
        "整体流程：作业一语料统计 → 作业二词典扩展与新词评估 → 作业三 LSI/LDA 主题建模 → "
        "文档嵌入 PCA/t-SNE/UMAP 可视化 → KNN 相似文档推荐 → 输出图表与 JSON 报告 → 生成 Word 文档。",
    )

    # 三、运行结果
    add_h1(doc, "三、运行结果与测试")
    add_h2(doc, "1、代码实现")
    add_body(doc, "新词爆发力评估：对作业二扩展词典中的新词计算 TF/DF 与语义新颖度。")
    add_code(
        doc,
        "newword_scores = evaluate_new_words(docs, seed_words, user_dict_words)\n"
        "# burst_index = 0.4*TF_norm + 0.3*DF_norm + 0.3*novelty",
    )
    add_body(doc, "主题模型与推荐：复用作业三 LSI 嵌入与 KNN 索引。")
    add_code(
        doc,
        "lsi_model, lsi_topics = train_lsi(tfidf, n_components=5)\n"
        "nn = build_knn_index(lsi_topics, metric='cosine', n_neighbors=6)\n"
        "recs = recommend(nn, lsi_topics, query_idx=0, k=5)",
    )

    add_h2(doc, "2、运行结果")
    add_body(doc, "（1）作业一至三数据联动概况：")
    add_body(
        doc,
        f"中文文档数：{integrated['n_chinese_docs']}，英文文档数：{integrated['n_english_docs']}\n"
        f"作业二 AI 词典：种子 {stats.get('AI 种子词数', '31')} 词 → 扩展 {stats.get('AI 最终词典词数', '223')} 词\n"
        f"前 5 维 SVD 能量占比（中文）：{summary['chinese_ai']['svd_energy']:.2%}，"
        f"（英文）：{summary['newsgroups']['svd_energy']:.2%}",
    )

    add_body(doc, "（2）作业一高频词与 TF-IDF 关键词（Top 5）：")
    hw1_lines = []
    for r in report["hw1_top_words"][:5]:
        hw1_lines.append(f"  {r['word']}：频次={r['freq']}")
    for r in report["hw1_tfidf"][:5]:
        hw1_lines.append(f"  {r['word']}：TF-IDF={r['weight']:.4f}")
    add_body(doc, "\n".join(hw1_lines))

    add_body(doc, "（3）新词爆发力指数 Top 10（课程设计自定义评估）：")
    add_body(doc, newword_lines(report["newword_scores"], 10))

    add_body(doc, "（4）LSI 主题词（中文语料）：")
    add_body(doc, topic_lines(zh["lsi_topics"]))

    add_body(doc, "（5）KNN 推荐示例（查询文档 #0，人工智能定义段落）：")
    rec_text = []
    for r in zh["knn_recommendations"]:
        rec_text.append(f"  推荐 #{r['doc_id']}（距离={r['distance']:.4f}）\n    {r['preview']}")
    add_body(doc, "\n".join(rec_text))

    add_h2(doc, "3、可视化结果")
    figures = [
        ("图 1：课程设计全流程示意图（作业一至三联动）", "pipeline_overview.png"),
        ("图 2：作业一语料高频词分布", "hw1_word_freq.png"),
        ("图 3：作业一 TF-IDF 关键词权重", "hw1_tfidf.png"),
        ("图 4：作业二 AI 词典规模变化", "hw2_dict_growth.png"),
        ("图 5：新词爆发力指数 Top 排名", "newword_burst_index.png"),
        ("图 6：SVD 奇异值谱（LSA/LSI 理论基础）", "chinese_ai_svd_spectrum.png"),
        ("图 7：LSI 主题高频词分布", "chinese_ai_lsi_topics.png"),
        ("图 8：LDA 主题高频词分布", "chinese_ai_lda_topics.png"),
        ("图 9：中文语料 LSI 嵌入 — PCA 可视化", "chinese_ai_embed_pca.png"),
        ("图 10：中文语料 LSI 嵌入 — t-SNE 可视化", "chinese_ai_embed_tsne.png"),
        ("图 11：中文语料 LSI 嵌入 — UMAP 可视化", "chinese_ai_embed_umap.png"),
        ("图 12：20 Newsgroups 语料 LSI 嵌入 — UMAP 对照", "newsgroups_embed_umap.png"),
    ]
    for caption, fname in figures:
        add_figure(doc, caption, os.path.join(OUTPUT_DIR, fname))

    # 四、总结
    add_h1(doc, "四、总结")
    add_body(
        doc,
        "本课程设计将作业一至三的实验成果整合为一条完整的人工智能文本挖掘流水线："
        "从语料采集与词频分析，到定制词典扩展与新词评估，再到 LSI/LDA 主题建模、"
        "文档嵌入可视化与 KNN 语义推荐。提出的「新词爆发力指数」结合 TF、DF 与语义新颖度，"
        "为作业二新词挖掘提供了可量化的评估视角。实验表明，三份作业数据在 AI 主题下具有良好的连贯性，"
        "降维后的文档嵌入能有效支撑相似文档推荐，PCA/t-SNE/UMAP 从不同角度揭示了文档聚类结构。"
        "后续可引入大语言模型嵌入（如 BERT）进一步提升语义表示质量。",
    )

    # 五、参考文献
    add_h1(doc, "五、参考文献")
    refs = [
        "[1] Gensim, Topic Modelling for humans, https://radimrehurek.com/gensim/, Accessed 2026-06-06.",
        "[2] Annoy (Approximate Nearest Neighbors Oh Yeah), https://github.com/spotify/annoy, Accessed 2026-06-06.",
        "[3] UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction, https://umap-learn.readthedocs.io/en/latest/.",
        "[4] scikit-learn User Guide — Decomposition and Manifold learning, https://scikit-learn.org/.",
        "[5] 结巴中文分词, https://github.com/fxsjy/jieba.",
        "[6] Mikolov et al., Efficient Estimation of Word Representations in Vector Space, arXiv:1301.3781, 2013.",
        "[7] Blei et al., Latent Dirichlet Allocation, JMLR 2003.",
        "[8] Matplotlib 官方文档, https://matplotlib.org/stable/.",
    ]
    for ref in refs:
        add_body(doc, ref, space_after=3)

    return doc


def _save_document(doc):
    """依次尝试多个路径，避免目标 docx 被 Word 打开时写入失败。"""
    candidates = [
        OUT_LOCAL,
        os.path.join(OUTPUT_DIR, "课程设计-陈喜盈(3223004641).docx"),
        os.path.join(BASE_DIR, "课程设计-陈喜盈(3223004641)_new.docx"),
    ]
    errors = []
    for path in candidates:
        try:
            doc.save(path)
            print("已保存:", path)
            return path
        except PermissionError as exc:
            errors.append(f"{path} ({exc})")
            print(f"[跳过] 文件被占用: {path}")
    raise PermissionError(
        "无法保存 Word 文档，请关闭已打开的同名校对文件后重试。\n" + "\n".join(errors)
    )


def main():
    if not os.path.isfile(REPORT_JSON):
        raise FileNotFoundError("请先运行 python main.py 生成 output 结果")

    _save_document(build_document())


if __name__ == "__main__":
    main()
